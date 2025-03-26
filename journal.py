import sys
import psycopg2
import configparser
import win32com.client as win32
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton, 
    QMessageBox, QCheckBox, QHBoxLayout, QFormLayout, QTableWidget, 
    QTableWidgetItem, QComboBox, QTextEdit, QTimeEdit, QDateEdit, QGroupBox, QListWidget, QHeaderView, QMenuBar, QAction, QDialog, QFileDialog, QSizePolicy, QSlider)
from PyQt5.QtCore import QTime, QDate, Qt, QTimer, QRegExp, QSettings
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt5.QtGui import QIcon, QPixmap, QFont, QRegExpValidator
import logging
import os
import resources
import pyqtgraph as pg
import numpy as np
import pyaudio
import wave
import speech_recognition as sr

# Определяем путь к директории, где находится исполняемый файл (.exe) или скрипт (.py)
if getattr(sys, 'frozen', False):  # Если запущен .exe файл
    current_dir = os.path.dirname(sys.executable)
else:
    current_dir = os.path.dirname(os.path.abspath(__file__))

# Настройка логирования, чтобы файл создавался в текущей директории
logging.basicConfig(
    filename=os.path.join(current_dir, 'app.log'),
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

logging.info("Приложение запущено.")

class LoginWindow(QWidget):
    def __init__(self):
        super().__init__()
        logging.info("Инициализация LoginWindow")
        self.initUI()
        self.load_credentials()
    def initUI(self):
        logging.info("Запуск initUI")
        self.setWindowTitle("Авторизация в shift_journal_db")

        main_layout = QVBoxLayout()
        form_layout = QFormLayout()

        self.user_input = QLineEdit()
        form_layout.addRow("Имя пользователя:", self.user_input)

        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.Password)
        form_layout.addRow("Пароль:", self.password_input)

        self.save_credentials_checkbox = QCheckBox("Сохранить данные")
        main_layout.addLayout(form_layout)
        main_layout.addWidget(self.save_credentials_checkbox)

        self.host_input = QLineEdit()
        form_layout.addRow("Хост:", self.host_input)

        button_layout = QHBoxLayout()
        self.login_button = QPushButton("Войти")
        self.login_button.clicked.connect(self.login)
        button_layout.addWidget(self.login_button)

        self.cancel_button = QPushButton("Отмена")
        self.cancel_button.clicked.connect(self.cancel)
        button_layout.addWidget(self.cancel_button)
        
        main_layout.addLayout(button_layout)
        self.setLayout(main_layout)
        logging.info("UI инициализирован")

    def load_credentials(self):
        logging.info("Загрузка учетных данных")
        config = configparser.ConfigParser()
        config.read("config.ini")
        
        if 'Credentials' in config:
            self.user_input.setText(config['Credentials'].get('username', ''))
            self.password_input.setText(config['Credentials'].get('password', ''))
            self.host_input.setText(config['Credentials'].get('host', ''))
            self.save_credentials_checkbox.setChecked(True)
            logging.info("Учетные данные загружены")
        else:
            logging.warning("Учетные данные не найдены в config.ini")

    def save_credentials(self):
        logging.info("Сохранение учетных данных")
        config = configparser.ConfigParser()
        config['Credentials'] = {
            'username': self.user_input.text(),
            'password': self.password_input.text(),
            'host': self.host_input.text()
        }
        with open("config.ini", "w") as config_file:
            config.write(config_file)
        logging.info("Учетные данные сохранены в config.ini")

    def login(self):
        host = self.host_input.text()
        user = self.user_input.text()
        password = self.password_input.text()

        logging.info("Попытка подключения к базе данных")
        try:
            self.connection = psycopg2.connect(
                host=host,
                user=user,
                password=password,
                dbname="shift_journal_db"
            )
            logging.info("Подключение к базе данных установлено")
            self.create_tables()
            QMessageBox.information(self, "Успешно", "Подключение к базе данных установлено")
            self.open_main_window()

            if self.save_credentials_checkbox.isChecked():
                self.save_credentials()
        except psycopg2.OperationalError as e:
            logging.error(f"Ошибка подключения к базе данных: {e}")
            QMessageBox.critical(self, "Ошибка", f"Не удалось подключиться к базе данных:\n{e}")

    def cancel(self):
        logging.info("Нажата кнопка 'Отмена', закрытие окна")
        self.close()

    def create_tables(self):
        cursor = self.connection.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS journal (
                id SERIAL PRIMARY KEY,
                date DATE,
                shift VARCHAR(50),
                time TIME,
                content TEXT,
                note TEXT
            )
        ''')

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS engineers (
                id SERIAL PRIMARY KEY,
                shift VARCHAR(50),
                date DATE,
                name VARCHAR(255)
            )
        ''')

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS engineers_info (
                id SERIAL PRIMARY KEY,
                full_name VARCHAR(255) NOT NULL,
                tab_number VARCHAR(50) NOT NULL
            )
        ''')

        self.connection.commit()

    def open_main_window(self):
        self.main_window = MainWindow(self.connection)
        self.main_window.show()
        self.close()

    def cancel(self):
        self.close()

class MainWindow(QWidget):
    def __init__(self, connection):
        super().__init__()
        self.connection = connection
        self.filters = {}  # Хранит текущие фильтры
        self.initUI()
                # Добавляем метод для проверки и восстановления соединения
        self.reconnect_if_needed()

        # Настройка таймера для периодической проверки соединения
        self.timer = QTimer()
        self.timer.timeout.connect(self.reconnect_if_needed)
        self.timer.start(300000)  # Проверять соединение каждые 5 минут

    def reconnect_if_needed(self):
        try:
            # Проверка состояния соединения
            self.connection.poll()
        except (psycopg2.InterfaceError, psycopg2.OperationalError):
            # Если соединение закрыто, создаем новое
            self.connection = psycopg2.connect(
                host=self.host_input.text(),
                user=self.user_input.text(),
                password=self.password_input.text(),
                dbname="shift_journal_db"
            )
            logging.info("Соединение с базой данных восстановлено.")
        # Подключаем обновление таблицы и списка инженеров к изменению смены и даты
        self.shift_combo.currentIndexChanged.connect(self.update_engineers_and_journal)
        self.date_edit.dateChanged.connect(self.update_engineers_and_journal)
        # Устанавливаем текущую смену и дату
        self.set_current_shift_and_date()
        
        # Подключаем обновление таблицы и списка инженеров к изменению смены и даты
        self.shift_combo.currentIndexChanged.connect(self.update_engineers_and_journal)
        self.date_edit.dateChanged.connect(self.update_engineers_and_journal)
        
        # Загружаем данные для текущей смены и даты
        self.update_engineers_and_journal()
        self.load_engineers()

    def set_current_shift_and_date(self):
        """Устанавливает текущую смену и дату в shift_combo и date_edit."""
        current_shift, shift_date = self.get_current_shift_and_date()
        
        # Устанавливаем текущую дату в виджет выбора даты
        self.date_edit.setDate(QDate.fromString(shift_date, "yyyy-MM-dd"))
        
        # Устанавливаем текущую смену в виджет выбора смены
        if current_shift == "1-я смена":
            self.shift_combo.setCurrentIndex(0)
        elif current_shift == "2-я смена":
            self.shift_combo.setCurrentIndex(1)

    def initUI(self):
        self.setWindowTitle("Журнал инженеров по АСУ")
        self.resize(1000, 700)  # Здесь можно указать нужные размеры в пикселях
        # Добавление меню бара
        self.menu_bar = QMenuBar(self)

        # Меню "Файл"
        file_menu = self.menu_bar.addMenu("Файл")
        self.engineers_list_action = QAction("Список инженеров", self)
        self.engineers_list_action.triggered.connect(self.open_engineers_list)
        self.export_to_word_action = QAction("Экспорт в Word", self)
                # Подключаем экспорт в Word
        self.export_to_word_action.triggered.connect(self.export_to_word)
        self.send_email_action = QAction("Отправить по почте", self)
        self.send_email_action.triggered.connect(self.send_email)
        file_menu.addAction(self.engineers_list_action)
        file_menu.addAction(self.export_to_word_action)
        file_menu.addAction(self.send_email_action)

        # Меню "Справка"
        help_menu = self.menu_bar.addMenu("Справка")
        self.about_action = QAction("О программе", self)
        self.about_action.triggered.connect(self.show_about_dialog)
        help_menu.addAction(self.about_action)
        # Добавление пункта меню для фильтров
        filters_menu = self.menu_bar.addMenu("Фильтры")
        self.open_filters_action = QAction("Настроить фильтры", self)
        self.open_filters_action.triggered.connect(self.open_filters_dialog)
        filters_menu.addAction(self.open_filters_action)

       # Основной компоновщик для окна
        main_layout = QVBoxLayout()
        main_layout.setMenuBar(self.menu_bar)

        content_layout = QHBoxLayout()

        # Левая часть с двумя GroupBox-ами (инженеры и данные для записи)
        left_layout = QVBoxLayout()

        # GroupBox для добавления инженеров на смену
        self.engineer_groupbox = QGroupBox("Добавление инженеров на смену")
        self.engineer_groupbox.setFixedWidth(300)  # Фиксированная ширина
        engineer_layout = QVBoxLayout()

        # ComboBox для выбора инженера
        self.engineer_select_combo = QComboBox()
        engineer_layout.addWidget(QLabel("Выберите инженера:"))
        engineer_layout.addWidget(self.engineer_select_combo)

# Поле для отображения добавленных инженеров
        self.engineer_list = QListWidget()  # Инициализация self.engineer_list
        self.engineer_list.setFixedHeight(100)  # Ограниченная высота списка инженеров
        self.engineer_list.setStyleSheet("font-size: 12px;")  # Увеличение размера шрифта

        engineer_layout.addWidget(QLabel("Инженеры на смену:"))
        engineer_layout.addWidget(self.engineer_list)


        # Кнопки добавить и удалить инженера
        engineer_button_layout = QHBoxLayout()
        self.add_engineer_button = QPushButton("Добавить инженера")
        self.add_engineer_button.clicked.connect(self.add_engineer)
        engineer_button_layout.addWidget(self.add_engineer_button)

        self.remove_engineer_button = QPushButton("Удалить инженера")
        self.remove_engineer_button.clicked.connect(self.remove_engineer)
        engineer_button_layout.addWidget(self.remove_engineer_button)

        engineer_layout.addLayout(engineer_button_layout)
        self.engineer_groupbox.setLayout(engineer_layout)

        # Добавляем GroupBox для инженеров на смену в левую часть
        left_layout.addWidget(self.engineer_groupbox)

        # GroupBox для полей ввода
        self.field_groupbox = QGroupBox("Данные для записи")
        self.field_groupbox.setFixedWidth(300)  # Фиксированная ширина
        field_layout = QVBoxLayout()

        from PyQt5.QtCore import QDate

        # Выпадающий календарь для выбора даты
        self.date_edit = QDateEdit()
        self.date_edit.setCalendarPopup(True)
        self.date_edit.setDate(QDate.currentDate())
        self.date_edit.setMaximumDate(QDate.currentDate())  # Запрет на выбор будущих дат

        field_layout.addWidget(QLabel("Выберите дату:"))
        field_layout.addWidget(self.date_edit)


        # Выбор смены
        self.shift_combo = QComboBox()
        self.shift_combo.addItems(["1-я смена", "2-я смена"])
        field_layout.addWidget(QLabel("Выберите смену:"))
        field_layout.addWidget(self.shift_combo)

           # Поле для ввода времени с автодобавлением двоеточия
        self.time_edit = QLineEdit()
        self.time_edit.setPlaceholderText("Введите время в формате ЧЧ:ММ")
        
        # Устанавливаем валидатор для проверки ввода времени и ограничения на 5 символов
        time_validator = QRegExpValidator(QRegExp(r"^(?:[01]\d|2[0-3]):[0-5]\d$"))
        self.time_edit.setValidator(time_validator)

        # Подключаем метод для автоматического добавления двоеточия
        self.time_edit.textChanged.connect(self.auto_insert_colon)
        
        field_layout.addWidget(QLabel("Время:"))
        field_layout.addWidget(self.time_edit)

        

               # Поле ввода содержания
        self.content_edit = QTextEdit()
        self.content_edit.setFixedHeight(265)

        # Устанавливаем увеличенный шрифт для поля ввода текста
        font = QFont("Arial", 12)  # Устанавливаем шрифт Arial и размер 14
        self.content_edit.setFont(font)  # Применяем увеличенный шрифт к полю

        # Кнопка записи голоса для поля "Содержание"
        self.voice_button_content = QPushButton("🎙 Записать голос")
        self.voice_button_content.setFixedHeight(40)  # Высота кнопки
        self.voice_button_content.clicked.connect(lambda: self.open_voice_recorder(self.content_edit))

        field_layout.addWidget(QLabel("Содержание:"))
        field_layout.addWidget(self.content_edit)
        field_layout.addWidget(self.voice_button_content)  # Добавляем кнопку записи голоса

        # Поле ввода примечаний
        self.note_edit = QTextEdit()
        self.note_edit.setFixedHeight(130)

        # Устанавливаем увеличенный шрифт для поля ввода примечаний
        font = QFont("Arial", 12)  # Устанавливаем шрифт Arial и размер 14
        self.note_edit.setFont(font)  # Применяем увеличенный шрифт к полю

        # Кнопка записи голоса для поля "Примечание"
        self.voice_button_note = QPushButton("🎙 Записать голос")
        self.voice_button_note.setFixedHeight(40)  # Высота кнопки
        self.voice_button_note.clicked.connect(lambda: self.open_voice_recorder(self.note_edit))

        field_layout.addWidget(QLabel("Примечание:"))
        field_layout.addWidget(self.note_edit)
        field_layout.addWidget(self.voice_button_note)  # Добавляем кнопку записи голоса

        # Кнопки добавить и удалить
        button_layout = QHBoxLayout()
        self.add_button = QPushButton("Добавить запись")
        self.add_button.clicked.connect(self.add_record)
        button_layout.addWidget(self.add_button)

        self.delete_button = QPushButton("Удалить запись")
        self.delete_button.clicked.connect(self.delete_record)
        button_layout.addWidget(self.delete_button)

        field_layout.addLayout(button_layout)
        self.field_groupbox.setLayout(field_layout)

        # Добавляем GroupBox с данными для записи в левую часть
        left_layout.addWidget(self.field_groupbox)
        left_layout.addStretch()  # Добавляем растяжение, чтобы не растягивались GroupBox-ы

        # Добавляем левую часть в основной компоновщик
        content_layout.addLayout(left_layout)

        # Таблица для отображения записей журнала в правой части
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["Дата", "Смена", "Время", "Содержание", "Примечание"])
        self.table.hideColumn(0)  # Скрыть колонку "Дата"

        # Настройка автоматической ширины столбцов и высоты строк
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

        # Растяжение для колонки "Содержание" (чтобы занимала оставшееся место)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)

        # Настройка политики размера для расширения таблицы
        self.table.setSizePolicy(self.table.sizePolicy().Expanding, self.table.sizePolicy().Expanding)

        content_layout.addWidget(self.table)
        main_layout.addLayout(content_layout)
        self.setLayout(main_layout)

        
        # Инициализация интерфейса завершена, теперь можно загрузить данные
        self.load_journal_data()
        self.load_engineers()

    def show_about_dialog(self):
        about_dialog = AboutDialog()
        about_dialog.exec_()

    def auto_insert_colon(self):
            """Автоматически добавляет двоеточие после ввода двух символов."""
            text = self.time_edit.text()
            
            # Если длина текста 2 и нет двоеточия, добавляем его
            if len(text) == 2 and ":" not in text:
                self.time_edit.setText(text + ":")    
    def validate_time(self):
        """Проверяет введенное время и обновляет стиль в зависимости от корректности."""
        time_text = self.time_edit.text()

        # Проверяем, что формат соответствует "HH:MM"
        if QRegExp(r"^(0[0-9]|1[0-9]|2[0-3]):([0-5][0-9])$").exactMatch(time_text):
            # Время введено корректно, зеленая рамка
            self.time_edit.setStyleSheet("border: 1px solid green;")
        else:
            # Некорректное время, красная рамка и сообщение
            self.time_edit.setStyleSheet("border: 1px solid red;")
            QMessageBox.warning(self, "Ошибка ввода", "Введите корректное время в формате ЧЧ:ММ.")
            self.time_edit.setText("")  # Очищаем поле при неверном вводе

    def open_filters_dialog(self):
        """Открывает окно фильтров."""
        self.filter_dialog = FilterDialog(self)
        self.filter_dialog.exec_()


    def apply_filters(self, filters):
        """Применяет фильтры к таблице журнала."""
        self.filters = filters
        self.load_journal_data()

    def reset_filters(self):
        """Сбрасывает фильтры и обновляет таблицу."""
        self.filters = {}
        self.load_journal_data()

    
    def send_email(self):
        """Отправляет записи журнала текущей смены на указанный адрес в Outlook с улучшенным форматированием."""
        
            # Преобразуем дату из виджета
        raw_date = self.date_edit.date().toPyDate()  # Получаем объект типа `date`

        # Формат для базы данных (yyyy-MM-dd)
        date_for_db = raw_date.strftime("%Y-%m-%d")

        # Формат для сообщения (dd-MM-yyyy)
        date_for_message = raw_date.strftime("%d-%m-%Y")

        try:
            shift = self.shift_combo.currentText()
            logging.info(f"Отправка отчета по смене {shift} за дату {date_for_message} начата.")

            # Загружаем инженеров на смене
            cursor = self.connection.cursor()
            cursor.execute("SELECT name FROM engineers WHERE date = %s AND shift = %s", (date_for_db, shift))
            engineers = cursor.fetchall()
            logging.info(f"Загружены инженеры на смене: {[engineer[0] for engineer in engineers]}")

            # Загружаем записи журнала с учетом сортировки времени для второй смены
            cursor.execute(
                """
                SELECT time, content, note 
                FROM journal 
                WHERE date = %s AND shift = %s 
                ORDER BY 
                    CASE 
                        WHEN shift = '2-я смена' AND time >= '00:00' AND time < '08:30' THEN date + INTERVAL '1 day'
                        ELSE date
                    END ASC,
                    time ASC
                """, 
                (date_for_db, shift)
            )
            journal_entries = cursor.fetchall()
            logging.info(f"Загружены записи журнала: {len(journal_entries)} записей.")

            # Формируем HTML-таблицу с датой для сообщения
            html_body = f"""
            <html>
            <head>
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                    }}
                    h2 {{
                        color: #0056b3;
                        border-bottom: 2px solid #0056b3;
                        padding-bottom: 5px;
                    }}
                    table {{
                        width: 100%;
                        border-collapse: collapse;
                        margin-top: 10px;
                    }}
                    th, td {{
                        padding: 8px;
                        border: 1px solid #ddd;
                    }}
                    th {{
                        background-color: #0056b3;
                        color: white;
                    }}
                </style>
            </head>
            <body>
                <h2>Смена: {shift}, Дата: {date_for_message}</h2>
                <h3>Инженеры на смене:</h3>
                <ul>
                    {''.join(f'<li>{engineer[0]}</li>' for engineer in engineers)}
                </ul>
                <h3>Записи журнала:</h3>
                <table>
                    <tr><th>Время</th><th>Содержание</th><th>Примечание</th></tr>
                    {''.join(f'<tr><td>{entry[0]}</td><td>{entry[1]}</td><td>{entry[2]}</td></tr>' for entry in journal_entries)}
                </table>
            </body>
            </html>
            """

            # Отправляем письмо
            outlook = win32.Dispatch("Outlook.Application")
            mail = outlook.CreateItem(0)
            mail.Subject = f"Сводка - {shift} от {date_for_message}"
            mail.HTMLBody = html_body
            mail.To = "v.ustimenko@bogatyr.kz"  # Замените на нужный адрес
            mail.CC = "asutp01@bogatyr.kz"  # Добавьте адрес в копию
            mail.Display()
            logging.info("Письмо успешно создано и открыто в Outlook.")

        except Exception as e:
            logging.error(f"Ошибка при отправке письма: {e}")
            QMessageBox.critical(self, "Ошибка отправки", f"Не удалось отправить сообщение: {e}")


    def format_journal_data(self):
        """Форматирует данные текущей смены в текст для отправки по почте"""
        data = ""
        for row in range(self.table.rowCount()):
            date = self.table.item(row, 0).text()
            shift = self.table.item(row, 1).text()
            time = self.table.item(row, 2).text()
            content = self.table.item(row, 3).text()
            note = self.table.item(row, 4).text()
            data += f"Дата: {date}, Смена: {shift}, Время: {time}\nСодержание: {content}\nПримечание: {note}\n\n"
        return data
    
    def export_to_word(self):
        """Экспортирует журнал и список инженеров в файл Word с оформлением."""

        # Получаем текущую дату и смену
        date = self.date_edit.date().toPyDate()
        formatted_date = date.strftime("%d %B %Y")
        shift = self.shift_combo.currentText()
        logging.info(f"Начат экспорт журнала для смены: {shift} на дату: {formatted_date}")

        try:
            # Запрашиваем инженеров на смене
            cursor = self.connection.cursor()
            cursor.execute("SELECT name FROM engineers WHERE date = %s AND shift = %s", (date, shift))
            engineers = cursor.fetchall()
            logging.info(f"Загружены инженеры на смене: {[engineer[0] for engineer in engineers]}")

            # Запрашиваем записи журнала
            cursor.execute("SELECT time, content, note FROM journal WHERE date = %s AND shift = %s", (date, shift))
            journal_entries = cursor.fetchall()
            logging.info(f"Загружены записи журнала: {len(journal_entries)} записей.")

            # Создаем новый документ Word
            doc = Document()
            title = doc.add_heading(f"Журнал смены: {shift}, Дата: {formatted_date}", level=1)
            title_run = title.runs[0]
            title_run.font.color.rgb = RGBColor(0, 85, 179)  # Синий цвет заголовка

            # Добавляем раздел с инженерами на смене
            engineers_heading = doc.add_heading("Инженеры на смене:", level=2)
            engineers_heading_run = engineers_heading.runs[0]
            engineers_heading_run.font.color.rgb = RGBColor(0, 85, 179)  # Синий цвет заголовка

            for engineer in engineers:
                doc.add_paragraph(engineer[0], style="List Bullet")

            # Создаем таблицу для записей журнала
            records_heading = doc.add_heading("Записи журнала:", level=2)
            records_heading_run = records_heading.runs[0]
            records_heading_run.font.color.rgb = RGBColor(0, 85, 179)  # Синий цвет заголовка

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            # Заголовок таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Время"
            hdr_cells[1].text = "Содержание"
            hdr_cells[2].text = "Примечание"

            # Применяем стиль заголовков таблицы
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 85, 179)  # Синий цвет заголовков

            # Добавляем записи журнала в таблицу
            for entry in journal_entries:
                row_cells = table.add_row().cells
                row_cells[0].text = entry[0].strftime("%H:%M")  # Время
                row_cells[1].text = entry[1]  # Содержание
                row_cells[2].text = entry[2]  # Примечание

            # Открываем диалоговое окно для сохранения файла
            options = QFileDialog.Options()
            file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить журнал как", f"Журнал_{shift}_{formatted_date}.docx", "Word Files (*.docx)", options=options)

            # Сохраняем файл
            if file_path:
                try:
                    doc.save(file_path)
                    QMessageBox.information(self, "Экспорт завершен", f"Журнал успешно экспортирован в файл {file_path}.")
                    logging.info(f"Журнал успешно экспортирован в файл: {file_path}")
                except Exception as e:
                    logging.error(f"Ошибка при сохранении файла: {e}")
                    QMessageBox.critical(self, "Ошибка экспорта", f"Не удалось сохранить файл: {e}")
            else:
                logging.info("Экспорт отменен пользователем.")

        except Exception as e:
            logging.error(f"Ошибка при экспорте журнала: {e}")
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при экспорте: {e}")




    def open_engineers_list(self):
        """Открывает окно списка инженеров."""
        logging.info("Открытие окна списка инженеров.")
        try:
            self.engineers_list_window = EngineersListWindow(self.connection, self)
            self.engineers_list_window.exec_()
            logging.info("Окно списка инженеров открыто.")
        except Exception as e:
            logging.error(f"Ошибка при открытии окна списка инженеров: {e}")

        # Подключаем обновление таблицы и списка инженеров к изменению смены и даты
        self.shift_combo.currentIndexChanged.connect(self.update_engineers_and_journal)
        self.date_edit.dateChanged.connect(self.update_engineers_and_journal)

    def update_engineers_and_journal(self):
        """Обновляет список инженеров и журнал."""
        logging.info("Обновление инженеров и журнала.")
        try:
            self.load_engineers()
            self.load_journal_data()
            logging.info("Инженеры и журнал успешно обновлены.")
        except Exception as e:
            logging.error(f"Ошибка при обновлении инженеров и журнала: {e}")

    def load_engineers(self):
        """Загружает список инженеров в ComboBox и обновляет QListWidget для выбранной смены и даты."""
        logging.info("Загрузка списка инженеров.")
        try:
            self.engineer_select_combo.clear()
            self.engineer_list.clear()

            # Получаем выбранную дату и смену
            date = self.date_edit.date().toString("yyyy-MM-dd")
            shift = self.shift_combo.currentText()
            logging.info(f"Дата: {date}, Смена: {shift}")

            cursor = self.connection.cursor()

            # Загружаем всех инженеров для добавления в смену
            cursor.execute("SELECT full_name FROM engineers_info")
            all_engineers = cursor.fetchall()
            for engineer in all_engineers:
                self.engineer_select_combo.addItem(engineer[0])
            logging.info("Все инженеры успешно загружены в ComboBox.")

            # Загружаем инженеров, уже добавленных на текущую смену и дату
            cursor.execute(
                "SELECT name FROM engineers WHERE date = %s AND shift = %s",
                (date, shift)
            )
            engineers = cursor.fetchall()
            for engineer in engineers:
                self.engineer_list.addItem(engineer[0])
            logging.info("Инженеры для текущей смены и даты успешно загружены.")
        except Exception as e:
            logging.error(f"Ошибка при загрузке списка инженеров: {e}")


    def add_engineer(self):
        """Добавляет выбранного инженера в список для текущей смены и даты."""
        engineer = self.engineer_select_combo.currentText()
        current_shift, shift_date = self.get_current_shift_and_date()
        logging.info(f"Попытка добавить инженера: {engineer}, Смена: {current_shift}, Дата: {shift_date}")

        # Проверяем, что текущая смена совпадает с выбранной
        if self.shift_combo.currentText() != current_shift or self.date_edit.date().toString("yyyy-MM-dd") != shift_date:
            logging.warning("Ошибка добавления инженера: несоответствие текущей смены и даты")
            QMessageBox.warning(self, "Ошибка", "Инженера можно добавлять только за текущую смену!")
            return

        if engineer and engineer not in [self.engineer_list.item(i).text() for i in range(self.engineer_list.count())]:
            self.engineer_list.addItem(engineer)
            
            try:
                cursor = self.connection.cursor()
                cursor.execute(
                    "INSERT INTO engineers (shift, date, name) VALUES (%s, %s, %s)",
                    (current_shift, shift_date, engineer)
                )
                self.connection.commit()
                logging.info(f"Инженер {engineer} успешно добавлен в смену {current_shift} на дату {shift_date}.")
            except Exception as e:
                logging.error(f"Ошибка при добавлении инженера {engineer} в базу данных: {e}")
        else:
            logging.info(f"Инженер {engineer} уже добавлен в список.")

    def remove_engineer(self):
        """Удаляет выбранного инженера из списка для текущей смены и даты."""
        current_shift, shift_date = self.get_current_shift_and_date()
        logging.info(f"Попытка удалить инженера из смены: {current_shift}, Дата: {shift_date}")

        # Проверяем, что выбранная смена и дата совпадают с текущими
        if (self.shift_combo.currentText() != current_shift or 
            self.date_edit.date().toString("yyyy-MM-dd") != shift_date):
            logging.warning("Ошибка удаления инженера: несоответствие текущей смены и даты")
            QMessageBox.warning(self, "Ошибка", "Инженера можно удалять только за текущую смену!")
            return

        # Удаление выбранного инженера из списка
        for item in self.engineer_list.selectedItems():
            engineer = item.text()
            logging.info(f"Удаление инженера {engineer} из смены {current_shift} на дату {shift_date}")

            try:
                # Удаляем инженера из базы данных только для текущей смены и даты
                cursor = self.connection.cursor()
                cursor.execute(
                    "DELETE FROM engineers WHERE shift=%s AND date=%s AND name=%s",
                    (current_shift, shift_date, engineer)
                )
                self.connection.commit()
                logging.info(f"Инженер {engineer} успешно удален из базы данных.")
                
                # Удаляем инженера из QListWidget
                self.engineer_list.takeItem(self.engineer_list.row(item))
                logging.info(f"Инженер {engineer} успешно удален из списка в интерфейсе.")
            except Exception as e:
                logging.error(f"Ошибка при удалении инженера {engineer} из базы данных: {e}")



    def get_current_shift_and_date(self):
        """Определяет текущую смену и соответствующую ей дату."""
        current_time = QTime.currentTime()
        current_date = QDate.currentDate()

        shift_start_1 = QTime(8, 30)
        shift_end_1 = QTime(20, 30)
        shift_start_2 = QTime(20, 30)
        shift_end_2 = QTime(8, 30)

        # Определяем смену и дату начала смены
        if shift_start_1 <= current_time < shift_end_1:
            return "1-я смена", current_date.toString("yyyy-MM-dd")
        elif current_time >= shift_start_2 or current_time < shift_end_2:
            # Вторая смена, для времени после полуночи - дата предыдущего дня
            shift_date = current_date.addDays(-1) if current_time < shift_end_2 else current_date
            return "2-я смена", shift_date.toString("yyyy-MM-dd")
        return None, None

    def add_record(self):
        self.reconnect_if_needed()  # Проверка соединения перед добавлением записи
        """Добавляет запись в журнал только за текущую смену и дату начала этой смены."""

        logging.info("Начало выполнения add_record")

        # Получаем текущую смену и её начальную дату
        current_shift, shift_date = self.get_current_shift_and_date()
        if current_shift is None or shift_date is None:
            logging.warning("Не удалось определить текущую смену или дату начала смены.")
            QMessageBox.warning(self, "Ошибка", "Не удалось определить текущую смену.")
            return
        logging.info(f"Текущая смена: {current_shift}, Дата начала смены: {shift_date}")

        # Проверяем, что выбранная смена и дата соответствуют текущим
        selected_date = self.date_edit.date().toString("yyyy-MM-dd")
        selected_shift = self.shift_combo.currentText()
        logging.info(f"Выбранная дата: {selected_date}, Выбранная смена: {selected_shift}")

        if selected_date != shift_date or selected_shift != current_shift:
            logging.warning("Попытка добавить запись за несоответствующую дату или смену.")
            QMessageBox.warning(self, "Ошибка", "Запись можно добавлять только за текущую дату и смену!")
            return

        # Устанавливаем дату начала смены для записи
        date = shift_date
        shift = current_shift
        time = self.time_edit.text().strip()  # Получение времени как текста и удаление лишних пробелов
        content = self.content_edit.toPlainText().strip()
        note = self.note_edit.toPlainText().strip()

        # Проверка обязательного поля "время"
        if not time:
            QMessageBox.warning(self, "Ошибка", "Поле 'Время' обязательно для заполнения.")
            return

        logging.info(f"Дата: {date}, Смена: {shift}, Время: {time}, Контент: {content}, Примечание: {note}")

        # Вставка записи в базу данных
        try:
            cursor = self.connection.cursor()
            cursor.execute(
                "INSERT INTO journal (date, shift, time, content, note) VALUES (%s, %s, %s, %s, %s)",
                (date, shift, time, content, note)
            )
            self.connection.commit()
            logging.info("Запись успешно добавлена в базу данных.")
        except Exception as e:
            self.connection.rollback()  # Выполняем откат транзакции при ошибке
            logging.error(f"Ошибка при добавлении записи в базу данных: {e}")
            QMessageBox.warning(self, "Ошибка", f"Не удалось добавить запись: {e}")
            return

        # Обновляем данные в таблице и очищаем поля
        self.load_journal_data()
        self.clear_input_fields()  # Очистка полей ввода
        QMessageBox.information(self, "Успешно", "Запись добавлена!")
        logging.info("Функция add_record завершена успешно.")


    
    def clear_input_fields(self):
        """Очищает поля ввода времени, содержания и примечания."""
        self.time_edit.clear()
        self.content_edit.clear()
        self.note_edit.clear()

    def load_journal_data(self):
        """Загружает данные из таблицы journal и разрешает редактирование полей 'Время', 'Содержание' и 'Примечание' только для текущей смены и даты при двойном нажатии."""

        # Отключаем обработчик изменений временно, чтобы избежать случайного сохранения при загрузке
        self.table.blockSignals(True)
        self.table.clearContents()

        # Получаем текущую смену и её дату
        current_shift, current_date = self.get_current_shift_and_date()

        # Получаем выбранную дату и смену
        date = self.date_edit.date().toString("yyyy-MM-dd")
        shift = self.shift_combo.currentText()

        cursor = self.connection.cursor()
        cursor.execute(
            """
            SELECT id, date, shift, time, content, note 
            FROM journal 
            WHERE date = %s AND shift = %s 
            ORDER BY 
                CASE 
                    WHEN shift = '2-я смена' AND time >= '00:00' AND time < '08:30' THEN date + INTERVAL '1 day'
                    ELSE date
                END ASC,
                time ASC
            """,
            (date, shift)
        )
        records = cursor.fetchall()

        # Устанавливаем количество строк в таблице
        self.table.setRowCount(len(records))
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["Дата", "Смена", "Время", "Содержание", "Примечание"])

        # Заполняем таблицу данными
        for row_index, row_data in enumerate(records):
            record_id, record_date, record_shift, record_time, record_content, record_note = row_data

            # Устанавливаем данные в ячейки и добавляем ID записи как пользовательские данные
            for column_index, item in enumerate([record_date, record_shift, record_time, record_content, record_note]):
                cell = QTableWidgetItem(str(item))
                if column_index == 0:
                    cell.setData(Qt.UserRole, record_id)  # Сохраняем ID записи в ячейке даты
                self.table.setItem(row_index, column_index, cell)

                # Разрешаем редактирование только для текущей смены и даты, и для полей "Время", "Содержание", "Примечание"
                if record_date == current_date and record_shift == current_shift and column_index in [2, 3, 4]:
                    cell.setFlags(cell.flags() | Qt.ItemIsEditable)  # Разрешаем редактирование
                else:
                    cell.setFlags(cell.flags() & ~Qt.ItemIsEditable)  # Остальные ячейки только для чтения

        # Подключаем обработчики событий
        self.table.cellDoubleClicked.connect(self.enable_editing)
        self.table.cellChanged.connect(self.update_record)

        # Включаем обработчик изменений
        self.table.blockSignals(False)


    def enable_editing(self, row, column):
        """Разрешает редактирование ячейки только при двойном клике, если это текущая смена и дата."""
        # Получаем текущую смену и её дату
        current_shift, current_date = self.get_current_shift_and_date()
        
        # Получаем данные строки
        record_date = self.table.item(row, 0).text()
        record_shift = self.table.item(row, 1).text()

        # Проверка: разрешено ли редактирование для этой строки и столбца
        if record_date == current_date and record_shift == current_shift and column in [2, 3, 4]:
            item = self.table.item(row, column)
            if item:
                item.setFlags(item.flags() | Qt.ItemIsEditable)

    def update_record(self, row, column):
        """Сохраняет изменения в базе данных при завершении редактирования ячейки, если это текущая смена и дата."""
        # Получаем текущую смену и её дату
        current_shift, current_date = self.get_current_shift_and_date()
        
        # Получаем данные строки
        record_date = self.table.item(row, 0).text()
        record_shift = self.table.item(row, 1).text()

        # Разрешаем сохранение только для текущей смены и даты
        if record_date != current_date or record_shift != current_shift:
            return

        # Получаем ID записи
        record_id = self.table.item(row, 0).data(Qt.UserRole)
        if not record_id:
            return

        # Получаем новое значение ячейки
        new_value = self.table.item(row, column).text()

        # Определяем, какое поле обновлять в базе данных
        field_map = {2: "time", 3: "content", 4: "note"}
        field = field_map.get(column)
        
        if not field:
            return

        # Обновляем запись в базе данных
        cursor = self.connection.cursor()
        cursor.execute(f"UPDATE journal SET {field} = %s WHERE id = %s", (new_value, record_id))
        self.connection.commit()





    def delete_record(self):
        """Удаляет выбранную запись из таблицы journal, если она относится к текущей дате и смене."""

        logging.info("Начало выполнения delete_record")

        # Получаем текущую смену и её начальную дату
        current_shift, shift_date = self.get_current_shift_and_date()
        if current_shift is None or shift_date is None:
            logging.warning("Не удалось определить текущую смену или дату начала смены.")
            QMessageBox.warning(self, "Ошибка", "Не удалось определить текущую смену.")
            return
        logging.info(f"Текущая смена: {current_shift}, Дата начала смены: {shift_date}")

        # Проверяем, что выбрана строка для удаления
        selected_row = self.table.currentRow()
        if selected_row == -1:
            logging.warning("Попытка удалить запись без выбора строки.")
            QMessageBox.warning(self, "Внимание", "Пожалуйста, выберите запись для удаления.")
            return

        # Получаем ID записи, дату, смену и время
        record_id = self.table.item(selected_row, 0).data(Qt.UserRole)  # ID записи сохранен в UserRole
        date = self.table.item(selected_row, 0).text()
        shift = self.table.item(selected_row, 1).text()
        time = self.table.item(selected_row, 2).text()
        logging.info(f"ID записи: {record_id}, Дата: {date}, Смена: {shift}, Время: {time}")

        # Проверяем, что запись относится к текущей дате и смене
        if date != shift_date or shift != current_shift:
            logging.warning("Попытка удалить запись за несоответствующую дату или смену.")
            QMessageBox.warning(self, "Ошибка", "Удалить можно только записи за текущую дату и смену!")
            return

        # Создаем диалог подтверждения с кнопками "Да" и "Нет"
        confirmation_dialog = QMessageBox(self)
        confirmation_dialog.setWindowTitle("Подтверждение удаления")
        confirmation_dialog.setText("Вы уверены, что хотите удалить выбранную запись?")
        confirmation_dialog.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        confirmation_dialog.button(QMessageBox.Yes).setText("Да")
        confirmation_dialog.button(QMessageBox.No).setText("Нет")
        confirmation_dialog.setDefaultButton(QMessageBox.No)

        # Ожидаем ответа от пользователя
        reply = confirmation_dialog.exec()

        # Если пользователь выбрал "Нет", отменяем операцию
        if reply != QMessageBox.Yes:
            logging.info("Удаление записи отменено пользователем.")
            return

        # Удаление записи из базы данных с использованием ID
        try:
            cursor = self.connection.cursor()
            cursor.execute("DELETE FROM journal WHERE id = %s", (record_id,))
            self.connection.commit()
            logging.info("Запись успешно удалена из базы данных.")
        except Exception as e:
            logging.error(f"Ошибка при удалении записи из базы данных: {e}")
            QMessageBox.warning(self, "Ошибка", f"Не удалось удалить запись: {e}")
            return

        # Обновляем данные таблицы после удаления
        self.load_journal_data()
        QMessageBox.information(self, "Успешно", "Запись удалена!")
        logging.info("Функция delete_record завершена успешно.")





    def closeEvent(self, event):
        if self.connection:
            self.connection.close()
    
    def open_voice_recorder(self, target_text_edit):
        """Открывает окно для записи голоса и преобразования в текст."""
        self.voice_recorder_dialog = VoiceRecorderDialog(target_text_edit)
        self.voice_recorder_dialog.exec_()

class EngineersListWindow(QDialog):
    def __init__(self, connection, main_window):
        super().__init__()
        self.connection = connection
        self.main_window = main_window  # Сохраняем ссылку на главное окно
        self.initUI()


    def initUI(self):
        self.setWindowTitle("Список инженеров")
        self.setGeometry(100, 100, 400, 300)

        layout = QVBoxLayout()

        # Таблица для отображения инженеров
        self.engineers_table = QTableWidget()
        self.engineers_table.setColumnCount(2)
        self.engineers_table.setHorizontalHeaderLabels(["ФИО", "Табельный номер"])
        layout.addWidget(self.engineers_table)

        # Поля для добавления нового инженера
        self.full_name_input = QLineEdit()
        layout.addWidget(QLabel("ФИО инженера:"))
        layout.addWidget(self.full_name_input)

        self.tab_number_input = QLineEdit()
        layout.addWidget(QLabel("Табельный номер:"))
        layout.addWidget(self.tab_number_input)

        # Кнопки добавить и удалить инженера
        button_layout = QHBoxLayout()
        self.add_engineer_button = QPushButton("Добавить инженера")
        self.add_engineer_button.clicked.connect(self.add_engineer)
        button_layout.addWidget(self.add_engineer_button)

        self.delete_engineer_button = QPushButton("Удалить инженера")
        self.delete_engineer_button.clicked.connect(self.delete_engineer)
        button_layout.addWidget(self.delete_engineer_button)

        layout.addLayout(button_layout)
        self.setLayout(layout)

        # Загрузка данных инженеров
        self.load_engineers_data()

    def load_engineers_data(self):
        cursor = self.connection.cursor()
        cursor.execute("SELECT full_name, tab_number FROM engineers_info")
        records = cursor.fetchall()

        # Обновляем таблицу
        self.engineers_table.setRowCount(len(records))
        for row_index, row_data in enumerate(records):
            for column_index, item in enumerate(row_data):
                self.engineers_table.setItem(row_index, column_index, QTableWidgetItem(str(item)))


    def add_engineer(self):
        logging.info("Начало выполнения add_engineer")

        full_name = self.full_name_input.text().strip()
        tab_number = self.tab_number_input.text().strip()
        logging.info(f"Введенное имя: {full_name}, Табельный номер: {tab_number}")

        if full_name and tab_number:
            try:
                cursor = self.connection.cursor()
                cursor.execute(
                    "INSERT INTO engineers_info (full_name, tab_number) VALUES (%s, %s)",
                    (full_name, tab_number)
                )
                self.connection.commit()
                logging.info(f"Инженер {full_name} с табельным номером {tab_number} добавлен в базу данных.")
            except Exception as e:
                logging.error(f"Ошибка при добавлении инженера в базу данных: {e}")
                QMessageBox.warning(self, "Ошибка", f"Не удалось добавить инженера: {e}")
                return

            # Очистить поля ввода
            self.full_name_input.clear()
            self.tab_number_input.clear()
            logging.info("Поля ввода очищены.")

            # Вызвать load_engineers_data для обновления таблицы в этом окне
            self.load_engineers_data()
            logging.info("Таблица с данными инженеров обновлена.")

            # Вызвать load_engineers в основном окне для обновления ComboBox
            self.main_window.load_engineers()
            logging.info("ComboBox в основном окне обновлен.")

            QMessageBox.information(self, "Успешно", "Инженер добавлен.")
            logging.info("Функция add_engineer завершена успешно.")
        else:
            logging.warning("Пользователь не заполнил все поля для добавления инженера.")
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, заполните все поля.")

    def delete_engineer(self):
        logging.info("Начало выполнения delete_engineer")

        selected_row = self.engineers_table.currentRow()
        if selected_row == -1:
            logging.warning("Попытка удалить инженера без выбора строки.")
            QMessageBox.warning(self, "Внимание", "Пожалуйста, выберите инженера для удаления.")
            return

        # Получаем данные выбранного инженера
        full_name = self.engineers_table.item(selected_row, 0).text()
        tab_number = self.engineers_table.item(selected_row, 1).text()
        logging.info(f"Удаление инженера: Имя - {full_name}, Табельный номер - {tab_number}")

        # Удаление записи из базы данных
        try:
            cursor = self.connection.cursor()
            cursor.execute(
                "DELETE FROM engineers_info WHERE full_name=%s AND tab_number=%s",
                (full_name, tab_number)
            )
            self.connection.commit()
            logging.info(f"Инженер {full_name} с табельным номером {tab_number} удален из базы данных.")
        except Exception as e:
            logging.error(f"Ошибка при удалении инженера из базы данных: {e}")
            QMessageBox.warning(self, "Ошибка", f"Не удалось удалить инженера: {e}")
            return

        # Обновление данных инженеров после удаления
        self.load_engineers_data()
        logging.info("Таблица с данными инженеров обновлена после удаления.")
        logging.info("Функция delete_engineer завершена успешно.")

class AboutDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("О программе")
        self.setFixedSize(450, 300)
        
        # Layout
        layout = QVBoxLayout()
        
        # Логотип
        logo_label = QLabel()
        pixmap = QPixmap(":/images/big-catalog-16444854041.jpg")  # Замените на путь к вашему логотипу
        logo_label.setPixmap(pixmap)
        logo_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(logo_label)
        
        # Название программы
        name_label = QLabel("Журнал инженеров по АСУ")
        name_label.setFont(QFont("Arial", 16, QFont.Bold))
        name_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(name_label)
        
        # Версия программы
        version_label = QLabel("Версия 1.1")
        version_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(version_label)
        
        # Контактные данные
        contact_label = QLabel("Отдел информационных технологий\nРИС ГТК\n+7 (7187) 22-38-13")
        contact_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(contact_label)
        
        self.setLayout(layout)

class FilterDialog(QDialog):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.page = 0  # Текущая страница
        self.page_size = 100  # Количество записей на странице
        self.filter_timer = QTimer(self)  # Таймер для дебаунсинга
        self.filter_timer.setSingleShot(True)  # Таймер однократного срабатывания
        self.filter_timer.timeout.connect(self.load_filtered_data)  # Связь с фильтрацией

        self.data_cache = []  # Кэш данных текущей страницы

        self.initUI()

    def initUI(self):
        self.setWindowTitle("Фильтр журнала")
        self.resize(1200, 700)

        layout = QVBoxLayout()

        # Поля для ввода содержания и примечания
        self.content_input = QLineEdit()
        self.content_input.textChanged.connect(self.on_filter_text_changed)
        layout.addWidget(QLabel("Содержание (ключевое слово):"))
        layout.addWidget(self.content_input)

        self.note_input = QLineEdit()
        self.note_input.textChanged.connect(self.on_filter_text_changed)
        layout.addWidget(QLabel("Примечание (ключевое слово):"))
        layout.addWidget(self.note_input)

        # Таблица для отображения результатов
        self.result_table = QTableWidget()
        self.result_table.setColumnCount(5)
        self.result_table.setHorizontalHeaderLabels(["Дата", "Смена", "Время", "Содержание", "Примечание"])
        self.result_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

        layout.addWidget(self.result_table)

        # Кнопки для пагинации
        pagination_layout = QHBoxLayout()
        self.prev_button = QPushButton("Назад")
        self.prev_button.clicked.connect(self.previous_page)
        self.prev_button.setEnabled(False)  # Изначально "Назад" отключена
        pagination_layout.addWidget(self.prev_button)

        self.next_button = QPushButton("Далее")
        self.next_button.clicked.connect(self.next_page)
        pagination_layout.addWidget(self.next_button)

        layout.addLayout(pagination_layout)
        self.setLayout(layout)

    def on_filter_text_changed(self):
        """Запускает таймер для дебаунсинга фильтрации."""
        self.filter_timer.start(300)

    def load_filtered_data(self):
        """Загружает данные с учетом фильтров и текущей страницы."""
        content_filter = self.content_input.text().strip()
        note_filter = self.note_input.text().strip()

        query = "SELECT date, shift, time, content, note FROM journal WHERE 1=1"
        params = []

        if content_filter:
            query += " AND content ILIKE %s"
            params.append(f"%{content_filter}%")

        if note_filter:
            query += " AND note ILIKE %s"
            params.append(f"%{note_filter}%")

        query += " ORDER BY date, time LIMIT %s OFFSET %s"
        params.extend([self.page_size, self.page * self.page_size])

        try:
            cursor = self.parent.connection.cursor()
            cursor.execute(query, params)
            self.data_cache = cursor.fetchall()  # Загружаем текущую страницу данных
            self.update_table()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить данные: {e}")

    def update_table(self):
        """Обновляет таблицу с данными текущей страницы."""
        self.result_table.setRowCount(len(self.data_cache))
        for row_index, row_data in enumerate(self.data_cache):
            for column_index, item in enumerate(row_data):
                cell = QTableWidgetItem(str(item))
                cell.setTextAlignment(Qt.AlignLeft | Qt.AlignTop)  # Выравнивание по левому верхнему углу
                cell.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)  # Только для чтения
                self.result_table.setItem(row_index, column_index, cell)

        # Включаем перенос текста для всех строк
        self.result_table.resizeRowsToContents()

        # Настраиваем ширины колонок
        self.set_column_widths()

        # Управляем состоянием кнопок пагинации
        self.prev_button.setEnabled(self.page > 0)  # Включаем кнопку "Назад", если это не первая страница
        self.next_button.setEnabled(len(self.data_cache) == self.page_size)  # Включаем "Далее", если есть ещё данные


    def set_column_widths(self):
        """Устанавливает ширины колонок."""
        # Устанавливаем фиксированный режим для "Дата", "Смена" и "Время"
        self.result_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Fixed)  # Дата
        self.result_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Fixed)  # Смена
        self.result_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Fixed)  # Время

        # Фиксируем ширины для "Дата", "Смена" и "Время"
        self.result_table.setColumnWidth(0, 120)  # Дата
        self.result_table.setColumnWidth(1, 100)  # Смена
        self.result_table.setColumnWidth(2, 100)  # Время

        # Устанавливаем минимальные и растягиваемые ширины для "Содержание" и "Примечание"
        self.result_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)  # Содержание
        self.result_table.setColumnWidth(3, 300)  # Устанавливаем стандартную ширину для "Содержание"

        self.result_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.Interactive)  # Примечание
        self.result_table.setColumnWidth(4, 200)  # Уменьшаем стандартную ширину для "Примечание"





    def next_page(self):
        """Переход к следующей странице."""
        self.page += 1
        self.load_filtered_data()

    def previous_page(self):
        """Переход к предыдущей странице."""
        if self.page > 0:  # Проверяем, что текущая страница не первая
            self.page -= 1  # Переход на предыдущую страницу
            self.load_filtered_data()
        else:
            self.prev_button.setEnabled(False)  # Отключаем кнопку "Назад", если это первая страница

class VoiceRecorderDialog(QDialog):
    def __init__(self, target_text_edit):
        super().__init__()
        self.target_text_edit = target_text_edit

        # Настройки для сохранения положения ползунка
        self.settings = QSettings("MyApp", "VoiceRecorder")

        self.initUI()

    def initUI(self):
        self.setWindowTitle("Запись голоса")
        self.resize(600, 500)

        layout = QVBoxLayout()

        # Визуализация
        self.plot_widget = pg.PlotWidget()
        self.plot_widget.setYRange(-32768, 32767)  # Диапазон амплитуды для 16-битного звука
        self.plot_curve = self.plot_widget.plot(pen="g")
        layout.addWidget(self.plot_widget)

        # Ползунок громкости
        self.volume_slider = QSlider(Qt.Horizontal)
        self.volume_slider.setMinimum(1)  # Минимальная громкость
        self.volume_slider.setMaximum(10)  # Максимальная громкость (10x)

        # Восстанавливаем сохранённое значение громкости
        saved_volume = self.settings.value("volume", 5, type=int)  # По умолчанию 5
        self.volume_slider.setValue(saved_volume)

        self.volume_slider.setTickInterval(1)
        self.volume_slider.setTickPosition(QSlider.TicksBelow)
        self.volume_slider.valueChanged.connect(self.update_volume_label)
        layout.addWidget(QLabel("Увеличение громкости:"))
        layout.addWidget(self.volume_slider)

        # Метка для текущей громкости
        self.volume_label = QLabel(f"Текущая громкость: {self.volume_slider.value()}x")
        layout.addWidget(self.volume_label)

        # Кнопка начала записи
        self.record_button = QPushButton("Начать запись")
        self.record_button.clicked.connect(self.start_recording)
        layout.addWidget(self.record_button)

        # Кнопка остановки записи
        self.stop_button = QPushButton("Остановить запись")
        self.stop_button.clicked.connect(self.stop_recording)
        self.stop_button.setEnabled(False)
        layout.addWidget(self.stop_button)

        # Информационный текст
        self.info_label = QLabel("Нажмите 'Начать запись' для записи голоса.")
        layout.addWidget(self.info_label)

        self.setLayout(layout)

        # Переменные для записи
        self.is_recording = False
        self.audio_file = "voice_recording.wav"
        self.timer = QTimer()
        self.volume_multiplier = self.volume_slider.value()  # Начальное увеличение громкости

    def update_volume_label(self):
        """Обновляет метку громкости при изменении значения ползунка."""
        self.volume_multiplier = self.volume_slider.value()
        self.volume_label.setText(f"Текущая громкость: {self.volume_multiplier}x")

    def start_recording(self):
        self.is_recording = True
        self.record_button.setEnabled(False)
        self.stop_button.setEnabled(True)
        self.info_label.setText("Идет запись...")

        self.audio = pyaudio.PyAudio()
        self.stream = self.audio.open(format=pyaudio.paInt16,
                                      channels=1,
                                      rate=16000,
                                      input=True,
                                      frames_per_buffer=1024)
        self.frames = []

        # Запуск таймера для обновления графика
        self.timer.timeout.connect(self.update_visualization)
        self.timer.start(50)

    def update_visualization(self):
        if self.is_recording:
            data = self.stream.read(1024, exception_on_overflow=False)
            audio_data = np.frombuffer(data, dtype=np.int16)

            # Применяем громкость (увеличиваем амплитуду сигнала)
            audio_data = np.clip(audio_data * self.volume_multiplier, -32768, 32767).astype(np.int16)
            self.frames.append(audio_data.tobytes())

            # Обновляем график
            self.plot_curve.setData(audio_data)

    def stop_recording(self):
        self.is_recording = False
        self.record_button.setEnabled(True)
        self.stop_button.setEnabled(False)
        self.info_label.setText("Запись завершена. Обработка...")

        self.timer.stop()
        self.stream.stop_stream()
        self.stream.close()
        self.audio.terminate()

        # Сохраняем запись
        with wave.open(self.audio_file, "wb") as wf:
            wf.setnchannels(1)
            wf.setsampwidth(self.audio.get_sample_size(pyaudio.paInt16))
            wf.setframerate(16000)
            wf.writeframes(b''.join(self.frames))

        # Преобразуем запись в текст
        self.process_audio()

    def process_audio(self):
        recognizer = sr.Recognizer()
        try:
            with sr.AudioFile(self.audio_file) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data, language="ru-RU")  # Используем русский язык
                self.target_text_edit.setText(text)
                self.info_label.setText("Голос успешно преобразован в текст.")
        except Exception as e:
            self.info_label.setText(f"Ошибка преобразования: {e}")

    def closeEvent(self, event):
        """Сохраняем положение ползунка громкости при закрытии окна."""
        self.settings.setValue("volume", self.volume_slider.value())
        event.accept()  # Закрываем окно

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon(":/images/bogatyr.ico"))  # Устанавливаем иконку для всего приложения
    window = LoginWindow()
    window.show()
    sys.exit(app.exec_())
